"""
Nova — an AI assistant hub.

Flask app with a menu launcher and tools (chat, summarize, jobs, scholarships),
user accounts, persisted chat history + saved summaries (SQLAlchemy — SQLite
locally, Postgres in production via DATABASE_URL), and subscription plans
(Free / Pro) with a pricing page.

The Gemini API key stays server-side (GEMINI_API_KEY); the browser only talks
to /api/chat.

Run locally:
    # put GEMINI_API_KEY in a .env next to this file, then:
    python app.py            # http://127.0.0.1:5004
"""
import json
import os
import re
import secrets
import urllib.error
import urllib.parse
import urllib.request
from datetime import datetime, timedelta, timezone
from functools import wraps
from pathlib import Path

from flask import (
    Flask, abort, flash, jsonify, redirect, render_template, request, session,
    url_for,
)
from flask_sqlalchemy import SQLAlchemy
from flask_wtf import CSRFProtect
from werkzeug.security import check_password_hash, generate_password_hash

try:
    from flask_limiter import Limiter
    from flask_limiter.util import get_remote_address
except ImportError:
    Limiter = None

BASE_DIR = Path(__file__).resolve().parent

try:
    from dotenv import load_dotenv
    load_dotenv(BASE_DIR / ".env")
except ImportError:
    pass

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx  # python-docx
except ImportError:
    docx = None

import base64
import io
try:
    import pandas as pd
    ANALYSIS_OK = True
except Exception:                    # noqa: BLE001
    pd = None
    ANALYSIS_OK = False

_plt_cache = {}


def _plt():
    """Lazy-load matplotlib (heavy) only when a chart is actually rendered."""
    if "plt" not in _plt_cache:
        import matplotlib
        matplotlib.use("Agg")        # headless server rendering
        import matplotlib.pyplot as plt
        _plt_cache["plt"] = plt
    return _plt_cache["plt"]


def _db_uri():
    url = os.environ.get("DATABASE_URL", "").strip()
    if not url:
        return "sqlite:///" + str(BASE_DIR / "nova.db")
    if url.startswith("postgres://"):        # Render/Heroku style -> SQLAlchemy style
        url = url.replace("postgres://", "postgresql://", 1)
    return url


app = Flask(__name__)
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", "dev-only-nova-key")
app.config["MAX_CONTENT_LENGTH"] = 2 * 1024 * 1024
app.config["SQLALCHEMY_DATABASE_URI"] = _db_uri()
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
# secure sessions (Secure cookie in production, i.e. when not in debug)
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=os.environ.get("FLASK_DEBUG", "1") != "1",
)

db = SQLAlchemy(app)
csrf = CSRFProtect(app)

if Limiter is not None:
    limiter = Limiter(key_func=get_remote_address, app=app,
                      default_limits=["300 per hour"],
                      storage_uri=os.environ.get("RATELIMIT_STORAGE_URI", "memory://"))
else:
    class _NoLimiter:
        def limit(self, *_a, **_k):
            return lambda f: f
    limiter = _NoLimiter()

GEMINI_MODEL = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash")
SYSTEM_PROMPT = (
    "You are Nova, a friendly, concise AI assistant. Answer helpfully and "
    "clearly. Use Markdown when useful — headings, bullet points, tables, and "
    "fenced code blocks with a language tag. If you're unsure, say so."
)
FREE_MAX_CONVERSATIONS = 5   # Pro removes this cap

# Expert personas — one-click system prompts for the chat.
PERSONAS = {
    "default": {"name": "Nova (general)", "system": SYSTEM_PROMPT},
    "resume": {"name": "Resume reviewer",
               "system": "You are an expert tech recruiter. Review resumes and answers "
                         "with specific, honest, actionable feedback. Use bullet points."},
    "study": {"name": "Study coach",
              "system": "You are a patient study coach. Explain clearly, check understanding, "
                        "and offer examples, analogies and practice questions."},
    "analyst": {"name": "Data analyst",
                "system": "You are a senior data analyst. Give precise, structured answers with "
                          "tables and formulas where useful, and note assumptions."},
    "coder": {"name": "Code helper",
              "system": "You are a senior software engineer. Give correct, idiomatic code in "
                        "fenced blocks with a language tag, then a short explanation."},
    "writer": {"name": "Writing assistant",
               "system": "You are a sharp writing assistant. Improve clarity and tone; offer "
                         "a polished version plus a note on what you changed."},
}


# --------------------------------------------------------------------------- #
# Models
# --------------------------------------------------------------------------- #
def utcnow():
    return datetime.now(timezone.utc)


class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(120), nullable=False)
    username = db.Column(db.String(80), unique=True, nullable=False)
    email = db.Column(db.String(200), unique=True)          # for payments + reset
    pwd_hash = db.Column(db.String(255), nullable=False)
    plan = db.Column(db.String(20), nullable=False, default="free")  # free | pro
    reset_token = db.Column(db.String(64))
    reset_expires = db.Column(db.String(40))
    created_at = db.Column(db.DateTime, default=utcnow)
    conversations = db.relationship("Conversation", backref="user", lazy=True,
                                    cascade="all, delete-orphan")
    summaries = db.relationship("SavedSummary", backref="user", lazy=True,
                                cascade="all, delete-orphan")

    @property
    def is_pro(self):
        return self.plan == "pro"


class Conversation(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    title = db.Column(db.String(160), default="New chat")
    is_public = db.Column(db.Boolean, default=False)
    share_token = db.Column(db.String(32), unique=True)
    created_at = db.Column(db.DateTime, default=utcnow)
    updated_at = db.Column(db.DateTime, default=utcnow)
    messages = db.relationship("Message", backref="conversation", lazy=True,
                               cascade="all, delete-orphan",
                               order_by="Message.id")


class Message(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    conversation_id = db.Column(db.Integer, db.ForeignKey("conversation.id"), nullable=False)
    role = db.Column(db.String(12), nullable=False)   # user | assistant
    content = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=utcnow)


class SavedSummary(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    source = db.Column(db.String(200), default="text")
    content = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=utcnow)


class Task(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    text = db.Column(db.String(300), nullable=False)
    done = db.Column(db.Boolean, default=False)
    due = db.Column(db.String(20))                 # optional YYYY-MM-DD
    created_at = db.Column(db.DateTime, default=utcnow)


class Note(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    title = db.Column(db.String(160), nullable=False)
    body = db.Column(db.Text, default="")
    created_at = db.Column(db.DateTime, default=utcnow)


class Document(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    name = db.Column(db.String(200), nullable=False)
    text = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=utcnow)


class Comment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    conversation_id = db.Column(db.Integer, db.ForeignKey("conversation.id"), nullable=False)
    author = db.Column(db.String(80), nullable=False, default="Guest")
    body = db.Column(db.String(1000), nullable=False)
    created_at = db.Column(db.DateTime, default=utcnow)


def ensure_schema():
    """Add newly-introduced columns to existing tables (no Alembic needed)."""
    from sqlalchemy import text, inspect
    with app.app_context():
        insp = inspect(db.engine)
        try:
            cols = {c["name"] for c in insp.get_columns("user")}
        except Exception:
            return
        wanted = {"email": "VARCHAR(200)", "reset_token": "VARCHAR(64)",
                  "reset_expires": "VARCHAR(40)"}
        for col, ddl in wanted.items():
            if col not in cols:
                try:
                    with db.engine.begin() as conn:
                        conn.execute(text('ALTER TABLE "user" ADD COLUMN %s %s' % (col, ddl)))
                except Exception as exc:  # noqa: BLE001
                    app.logger.warning("Could not add column %s: %s", col, exc)


with app.app_context():
    db.create_all()
ensure_schema()


# --------------------------------------------------------------------------- #
# Auth helpers
# --------------------------------------------------------------------------- #
def current_user():
    uid = session.get("user_id")
    return db.session.get(User, uid) if uid else None


def login_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if current_user() is None:
            flash("Please log in to use that feature.", "error")
            return redirect(url_for("login", next=request.path))
        return view(*args, **kwargs)
    return wrapped


def pro_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        u = current_user()
        if u is None:
            flash("Please log in first.", "error")
            return redirect(url_for("login", next=request.path))
        if not u.is_pro:
            flash("That's a Pro feature — upgrade to unlock it.", "error")
            return redirect(url_for("pricing"))
        return view(*args, **kwargs)
    return wrapped


@app.context_processor
def inject_user():
    return {"user": current_user()}


# --------------------------------------------------------------------------- #
# Hardening: security headers + error pages
# --------------------------------------------------------------------------- #
@app.after_request
def security_headers(resp):
    resp.headers["X-Content-Type-Options"] = "nosniff"
    resp.headers["X-Frame-Options"] = "DENY"
    resp.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    resp.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
    # allow the CDN libs used to render rich AI responses (markdown/code/math/diagrams)
    resp.headers["Content-Security-Policy"] = (
        "default-src 'self'; "
        "img-src 'self' data:; "
        "style-src 'self' 'unsafe-inline' https://fonts.googleapis.com https://cdn.jsdelivr.net; "
        "font-src 'self' https://fonts.gstatic.com https://cdn.jsdelivr.net; "
        "script-src 'self' https://cdn.jsdelivr.net; "
        "connect-src 'self'; base-uri 'self'; frame-ancestors 'none'"
    )
    return resp


@app.errorhandler(404)
def not_found(_e):
    return render_template("error.html", code=404, message="That page doesn't exist."), 404


@app.errorhandler(429)
def rate_limited(_e):
    if request.path.startswith("/api/"):
        return jsonify(reply="You're going too fast — please slow down.", error=True), 429
    return render_template("error.html", code=429,
                           message="Too many requests — please slow down."), 429


@app.errorhandler(500)
def server_error(_e):
    return render_template("error.html", code=500, message="Something went wrong."), 500


# --------------------------------------------------------------------------- #
# Content: menu + scholarships
# --------------------------------------------------------------------------- #
FEATURES = [
    {"icon": "💬", "title": "Chat with AI",
     "desc": "Ask anything, with rich formatted answers.", "endpoint": "chat_page", "gated": False, "pro": False},
    {"icon": "📝", "title": "Summarize text or document",
     "desc": "Paste text or upload Word/PDF/txt and get the gist.", "endpoint": "summarize", "gated": True, "pro": False},
    {"icon": "💼", "title": "Job search",
     "desc": "Find the latest remote jobs by keyword, category & location.", "endpoint": "jobs", "gated": True, "pro": False},
    {"icon": "🎓", "title": "Scholarship updates",
     "desc": "Browse current scholarship opportunities.", "endpoint": "scholarships", "gated": True, "pro": False},
    {"icon": "📎", "title": "Chat with your document",
     "desc": "Upload a PDF/Word/CSV and ask questions about it.", "endpoint": "docchat", "gated": True, "pro": False},
]

# Pro-only tools. `endpoint` = live; None = coming soon.
PRO_TOOLS = [
    {"icon": "📊", "title": "Data analysis", "endpoint": "analyze",
     "desc": "Upload a CSV/Excel — clean, profile, visualize & get AI insights."},
    {"icon": "📚", "title": "AI Study Mode", "endpoint": "study",
     "desc": "Turn notes into flashcards, quizzes and key points."},
    {"icon": "✅", "title": "Smart Productivity", "endpoint": "productivity",
     "desc": "Tasks, notes, and meeting-transcript action items."},
    {"icon": "⚡", "title": "Custom Workflows", "endpoint": "workflow",
     "desc": "Chain AI steps: summarize → translate → improve → save to notes."},
    {"icon": "🤝", "title": "Collaboration", "endpoint": None,
     "desc": "Share chats (live) and comment. Team co-editing coming soon."},
]

SCHOLARSHIPS = [
    {"title": "Mastercard Foundation Scholars Program", "provider": "Mastercard Foundation",
     "level": "Undergraduate & Master's", "region": "Africa",
     "blurb": "Full funding for academically talented young people, especially from Africa.",
     "url": "https://mastercardfdn.org/all/scholars/"},
    {"title": "Chevening Scholarships", "provider": "UK Government", "level": "Master's", "region": "Global",
     "blurb": "Fully-funded one-year master's study in the UK for future leaders.",
     "url": "https://www.chevening.org/"},
    {"title": "DAAD Scholarships", "provider": "DAAD (Germany)", "level": "Master's & PhD", "region": "Global",
     "blurb": "Funding to study and research in Germany.",
     "url": "https://www.daad.de/en/study-and-research-in-germany/scholarships/"},
    {"title": "Fulbright Foreign Student Program", "provider": "US Government", "level": "Master's & PhD", "region": "Global",
     "blurb": "Graduate study, research and teaching in the United States.",
     "url": "https://foreign.fulbrightonline.org/"},
    {"title": "Commonwealth Scholarships", "provider": "Commonwealth (UK)", "level": "Master's & PhD", "region": "Commonwealth",
     "blurb": "Funded UK study for students from Commonwealth nations.",
     "url": "https://cscuk.fcdo.gov.uk/scholarships/"},
    {"title": "Women Techmakers Scholarship", "provider": "Google", "level": "Undergraduate & Graduate", "region": "Global",
     "blurb": "Support for students in computer science and technology.",
     "url": "https://www.womentechmakers.com/scholars"},
]

JOB_CATEGORIES = [
    ("", "All categories"), ("software-dev", "Software Development"), ("data", "Data"),
    ("devops", "DevOps / Sysadmin"), ("design", "Design"), ("product", "Product"),
    ("marketing", "Marketing"), ("sales", "Sales / Business"),
    ("customer-support", "Customer Support"), ("finance-legal", "Finance / Legal"),
    ("hr", "HR"), ("qa", "QA"), ("writing", "Writing"), ("all-others", "All others"),
]

PLANS = {
    "free": {"name": "Free", "price": "₦0", "period": "forever",
             "features": ["Chat with AI", "Summarize documents", "Job search",
                          "Scholarship updates", f"Save up to {FREE_MAX_CONVERSATIONS} chats"]},
    "pro": {"name": "Pro", "price": "₦2,500", "period": "/month",
            "features": ["Everything in Free", "Unlimited saved chats & summaries",
                         "AI Study Mode", "Smart Productivity (tasks, reminders, notes)",
                         "Collaboration & sharing", "Custom automation workflows",
                         "Priority responses"]},
}


# --------------------------------------------------------------------------- #
# Gemini
# --------------------------------------------------------------------------- #
def gemini_reply(messages, system=None):
    key = os.environ.get("GEMINI_API_KEY")
    if not key:
        return None, "AI chat isn't configured yet — a GEMINI_API_KEY is needed."
    url = (f"https://generativelanguage.googleapis.com/v1beta/models/"
           f"{GEMINI_MODEL}:generateContent?key={key}")
    contents = []
    for m in messages:
        text = str(m.get("text", "")).strip()[:4000]
        if not text:
            continue
        role = "model" if m.get("role") == "assistant" else "user"
        contents.append({"role": role, "parts": [{"text": text}]})
    if not contents:
        return None, "Say something and I'll reply."
    body = json.dumps({
        "systemInstruction": {"parts": [{"text": system or SYSTEM_PROMPT}]},
        "contents": contents,
        "generationConfig": {"temperature": 0.7, "maxOutputTokens": 1200},
    }).encode("utf-8")
    req = urllib.request.Request(url, data=body,
                                 headers={"Content-Type": "application/json"}, method="POST")
    try:
        with urllib.request.urlopen(req, timeout=40) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        app.logger.warning("Gemini HTTP %s: %s", exc.code, exc.read().decode('utf-8', 'ignore')[:200])
        if exc.code in (400, 403):
            return None, "The AI key looks invalid or lacks access."
        if exc.code == 429:
            return None, "The AI is busy (quota/rate limit) — try again shortly."
        return None, f"AI service error ({exc.code})."
    except Exception as exc:  # noqa: BLE001
        app.logger.warning("Gemini error: %s", exc)
        return None, "Couldn't reach the AI service. Please try again."
    candidates = data.get("candidates") or []
    if not candidates:
        return None, "The AI didn't return a response (it may have been filtered)."
    parts = (candidates[0].get("content") or {}).get("parts") or []
    text = "".join(p.get("text", "") for p in parts).strip()
    return (text or "…(empty response)"), None


# --------------------------------------------------------------------------- #
# Summarizer + file extraction
# --------------------------------------------------------------------------- #
STOPWORDS = set("""a an the and or but if while of to in on for with as by at from
into is are was were be been being this that these those it its he she they we you
i me my our your their his her them us do does did have has had not no so than then
there here what which who whom how when where why can could should would will just
about above below over under again more most some such only own same too very""".split())


def summarize_text(text, max_sentences=5):
    text = re.sub(r"\s+", " ", text or "").strip()
    if not text:
        return "", {}
    sentences = [s.strip() for s in re.split(r"(?<=[.!?])\s+", text) if s.strip()]
    if len(sentences) <= max_sentences:
        n = len(text.split())
        return text, {"sentences": len(sentences), "kept": len(sentences), "words_in": n, "words_out": n}
    freq = {}
    for w in re.findall(r"[a-zA-Z']+", text.lower()):
        if w in STOPWORDS or len(w) <= 2:
            continue
        freq[w] = freq.get(w, 0) + 1
    if not freq:
        chosen = sentences[:max_sentences]
    else:
        scored = []
        for idx, s in enumerate(sentences):
            words = re.findall(r"[a-zA-Z']+", s.lower())
            if not words:
                continue
            scored.append((sum(freq.get(w, 0) for w in words) / (len(words) ** 0.5), idx, s))
        scored.sort(reverse=True)
        chosen = [s for _, _, s in sorted(scored[:max_sentences], key=lambda t: t[1])]
    summary = " ".join(chosen)
    return summary, {"sentences": len(sentences), "kept": len(chosen),
                     "words_in": len(text.split()), "words_out": len(summary.split())}


def extract_file_text(fs):
    name = (fs.filename or "").lower()
    if name.endswith((".txt", ".md", ".csv")):
        return fs.read().decode("utf-8", "ignore"), None
    if name.endswith(".pdf"):
        if PdfReader is None:
            return None, "PDF support isn't installed on the server."
        try:
            return "\n".join((p.extract_text() or "") for p in PdfReader(fs).pages), None
        except Exception as exc:  # noqa: BLE001
            return None, f"Couldn't read that PDF: {exc}"
    if name.endswith(".docx"):
        if docx is None:
            return None, "Word (.docx) support isn't installed on the server."
        try:
            return "\n".join(p.text for p in docx.Document(fs).paragraphs), None
        except Exception as exc:  # noqa: BLE001
            return None, f"Couldn't read that Word file: {exc}"
    if name.endswith(".doc"):
        return None, "Old .doc isn't supported — save as .docx or PDF."
    return None, "Unsupported file — upload .txt, .md, .pdf or .docx."


# --------------------------------------------------------------------------- #
# Jobs (Remotive)
# --------------------------------------------------------------------------- #
def fetch_jobs(search="", category="", company="", location="", fetch=50, show=24):
    url = "https://remotive.com/api/remote-jobs?limit=" + str(fetch)
    if search:
        url += "&search=" + urllib.parse.quote(search)
    if category:
        url += "&category=" + urllib.parse.quote(category)
    if company:
        url += "&company_name=" + urllib.parse.quote(company)
    req = urllib.request.Request(url, headers={
        "User-Agent": "Mozilla/5.0 (compatible; NovaJobs/1.0)", "Accept": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=20) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except Exception as exc:  # noqa: BLE001
        app.logger.warning("Remotive error: %s", exc)
        return None, "Couldn't fetch jobs right now. Please try again."
    loc = location.strip().lower()
    jobs = []
    for j in (data.get("jobs") or []):
        jl = (j.get("candidate_required_location") or "Remote")
        if loc and loc not in jl.lower() and "worldwide" not in jl.lower() and "anywhere" not in jl.lower():
            continue
        jobs.append({"title": j.get("title", ""), "company": j.get("company_name", ""),
                     "location": jl, "category": j.get("category", ""),
                     "date": (j.get("publication_date") or "")[:10], "url": j.get("url", "#")})
        if len(jobs) >= show:
            break
    return jobs, None


# --------------------------------------------------------------------------- #
# Routes — menu + auth
# --------------------------------------------------------------------------- #
@app.route("/")
def index():
    return render_template("index.html", features=FEATURES, pro_tools=PRO_TOOLS)


@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "GET":
        return render_template("register.html")
    f = request.form
    name, username = f.get("name", "").strip(), f.get("username", "").strip()
    email = f.get("email", "").strip().lower()
    pwd, confirm = f.get("password", ""), f.get("confirm", "")
    if not name:
        flash("Please enter your name.", "error")
    elif len(username) < 3:
        flash("Username must be at least 3 characters.", "error")
    elif not re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", email):
        flash("Please enter a valid email.", "error")
    elif len(pwd) < 6:
        flash("Password must be at least 6 characters.", "error")
    elif pwd != confirm:
        flash("Passwords do not match.", "error")
    elif User.query.filter_by(username=username).first():
        flash("That username is taken.", "error")
    elif User.query.filter_by(email=email).first():
        flash("That email is already registered.", "error")
    else:
        db.session.add(User(name=name, username=username, email=email,
                            pwd_hash=generate_password_hash(pwd)))
        db.session.commit()
        flash("Account created — please log in.", "success")
        return redirect(url_for("login"))
    return render_template("register.html", form=f), 400


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "GET":
        return render_template("login.html", next=request.args.get("next", ""))
    username = request.form.get("username", "").strip()
    pwd = request.form.get("password", "")
    nxt = request.form.get("next", "")
    u = User.query.filter_by(username=username).first()
    if u and check_password_hash(u.pwd_hash, pwd):
        session.clear()
        session["user_id"] = u.id
        flash(f"Welcome, {u.name}.", "success")
        return redirect(nxt if nxt and nxt.startswith("/") else url_for("index"))
    flash("Invalid username or password.", "error")
    return render_template("login.html", next=nxt), 400


@app.route("/logout")
def logout():
    session.clear()
    flash("Logged out.", "success")
    return redirect(url_for("index"))


def send_email(to, subject, html):
    key = os.environ.get("RESEND_API_KEY")
    sender = os.environ.get("MAIL_FROM", "Nova <onboarding@resend.dev>")
    if not key:
        return False
    body = json.dumps({"from": sender, "to": [to], "subject": subject,
                       "html": html}).encode("utf-8")
    req = urllib.request.Request("https://api.resend.com/emails", data=body,
        headers={"Authorization": "Bearer " + key, "Content-Type": "application/json"},
        method="POST")
    try:
        with urllib.request.urlopen(req, timeout=20) as resp:
            resp.read()
        return True
    except Exception as exc:  # noqa: BLE001
        app.logger.warning("Resend failed: %s", exc)
        return False


@app.route("/forgot", methods=["GET", "POST"])
@limiter.limit("10 per hour")
def forgot():
    if request.method == "GET":
        return render_template("forgot.html")
    email = request.form.get("email", "").strip().lower()
    user = User.query.filter_by(email=email).first() if email else None
    dev_link = None
    if user:
        token = secrets.token_urlsafe(24)
        user.reset_token = token
        user.reset_expires = (utcnow() + timedelta(hours=1)).isoformat()
        db.session.commit()
        link = url_for("reset_password", token=token, _external=True)
        sent = send_email(user.email, "Reset your Nova password",
                          "<p>Reset your password (valid 1 hour):</p>"
                          "<p><a href='%s'>%s</a></p>" % (link, link))
        if not sent:
            dev_link = link   # shown on-screen when email isn't configured yet
    flash("If that email is registered, we've sent a reset link.", "success")
    return render_template("forgot.html", dev_link=dev_link)


@app.route("/reset/<token>", methods=["GET", "POST"])
def reset_password(token):
    user = User.query.filter_by(reset_token=token).first()
    valid = False
    if user and user.reset_expires:
        try:
            valid = datetime.fromisoformat(user.reset_expires) > utcnow()
        except ValueError:
            valid = False
    if not valid:
        flash("That reset link is invalid or has expired.", "error")
        return redirect(url_for("forgot"))
    if request.method == "GET":
        return render_template("reset.html", token=token)
    pwd, confirm = request.form.get("password", ""), request.form.get("confirm", "")
    if len(pwd) < 6:
        flash("Password must be at least 6 characters.", "error")
        return render_template("reset.html", token=token), 400
    if pwd != confirm:
        flash("Passwords do not match.", "error")
        return render_template("reset.html", token=token), 400
    user.pwd_hash = generate_password_hash(pwd)
    user.reset_token = None
    user.reset_expires = None
    db.session.commit()
    flash("Password updated — please log in.", "success")
    return redirect(url_for("login"))


# --------------------------------------------------------------------------- #
# Subscription
# --------------------------------------------------------------------------- #
@app.route("/pricing")
def pricing():
    return render_template("pricing.html", plans=PLANS)


PRO_PRICE_KOBO = int(os.environ.get("PRO_PRICE_KOBO", "250000"))  # ₦2,500


@app.route("/subscribe", methods=["POST"])
@login_required
def subscribe():
    user = current_user()
    key = os.environ.get("PAYSTACK_SECRET_KEY")
    # Demo fallback when Paystack isn't configured (or user has no email)
    if not key or not user.email:
        user.plan = "pro"
        db.session.commit()
        flash("🎉 You're on Pro! (Demo upgrade — add a Paystack key for real billing.)", "success")
        return redirect(url_for("index"))
    # Real Paystack: initialize a transaction and redirect to checkout
    body = json.dumps({
        "email": user.email,
        "amount": PRO_PRICE_KOBO,
        "callback_url": url_for("payment_callback", _external=True),
        "metadata": {"user_id": user.id},
    }).encode("utf-8")
    req = urllib.request.Request(
        "https://api.paystack.co/transaction/initialize", data=body,
        headers={"Authorization": "Bearer " + key, "Content-Type": "application/json"},
        method="POST")
    try:
        with urllib.request.urlopen(req, timeout=20) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        auth_url = data["data"]["authorization_url"]
    except Exception as exc:  # noqa: BLE001
        app.logger.warning("Paystack init failed: %s", exc)
        flash("Couldn't start checkout — please try again.", "error")
        return redirect(url_for("pricing"))
    return redirect(auth_url)


@app.route("/payment/callback")
@login_required
def payment_callback():
    user = current_user()
    key = os.environ.get("PAYSTACK_SECRET_KEY")
    reference = request.args.get("reference", "")
    if not key or not reference:
        flash("Payment could not be verified.", "error")
        return redirect(url_for("pricing"))
    req = urllib.request.Request(
        "https://api.paystack.co/transaction/verify/" + urllib.parse.quote(reference),
        headers={"Authorization": "Bearer " + key})
    try:
        with urllib.request.urlopen(req, timeout=20) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        ok = data.get("data", {}).get("status") == "success"
    except Exception as exc:  # noqa: BLE001
        app.logger.warning("Paystack verify failed: %s", exc)
        ok = False
    if ok:
        user.plan = "pro"
        db.session.commit()
        flash("🎉 Payment confirmed — welcome to Pro!", "success")
        return redirect(url_for("index"))
    flash("Payment wasn't completed. You can try again.", "error")
    return redirect(url_for("pricing"))


@app.route("/downgrade", methods=["POST"])
@login_required
def downgrade():
    user = current_user()
    user.plan = "free"
    db.session.commit()
    flash("Switched back to the Free plan.", "success")
    return redirect(url_for("pricing"))


# --------------------------------------------------------------------------- #
# Chat + history
# --------------------------------------------------------------------------- #
@app.route("/chat")
def chat_page():
    convo, messages = None, []
    user = current_user()
    cid = request.args.get("c", type=int)
    if user and cid:
        convo = db.session.get(Conversation, cid)
        if not convo or convo.user_id != user.id:
            abort(404)
        messages = [{"role": m.role, "text": m.content} for m in convo.messages]
    personas = [(k, v["name"]) for k, v in PERSONAS.items()]
    return render_template("chat.html", conversation_id=convo.id if convo else None,
                           preload=messages, personas=personas)


@app.route("/api/chat", methods=["POST"])
@csrf.exempt
@limiter.limit("20 per minute")
def api_chat():
    payload = request.get_json(silent=True) or {}
    messages = payload.get("messages")
    if not isinstance(messages, list) or not messages:
        return jsonify(reply="No message received.", error=True), 400
    persona = PERSONAS.get(payload.get("persona", "default"), PERSONAS["default"])
    reply, err = gemini_reply(messages[-20:], system=persona["system"])
    if err:
        return jsonify(reply=err, error=True)

    convo_id = payload.get("conversation_id")
    user = current_user()
    if user:
        last_user = next((m for m in reversed(messages) if m.get("role") == "user"), None)
        convo = db.session.get(Conversation, convo_id) if convo_id else None
        if not convo or convo.user_id != user.id:
            # enforce free-plan cap on number of saved conversations
            if not user.is_pro:
                count = Conversation.query.filter_by(user_id=user.id).count()
                if count >= FREE_MAX_CONVERSATIONS:
                    return jsonify(reply=reply, conversation_id=None,
                                   notice="Free plan saves up to %d chats. Upgrade to Pro to keep more."
                                   % FREE_MAX_CONVERSATIONS)
            title = (last_user["text"][:60] if last_user else "New chat")
            convo = Conversation(user_id=user.id, title=title)
            db.session.add(convo)
            db.session.flush()
        if last_user:
            db.session.add(Message(conversation_id=convo.id, role="user", content=last_user["text"][:8000]))
        db.session.add(Message(conversation_id=convo.id, role="assistant", content=reply[:8000]))
        convo.updated_at = utcnow()
        db.session.commit()
        return jsonify(reply=reply, conversation_id=convo.id)
    return jsonify(reply=reply, conversation_id=None)


@app.route("/history")
@login_required
def history():
    user = current_user()
    convos = Conversation.query.filter_by(user_id=user.id).order_by(Conversation.updated_at.desc()).all()
    sums = SavedSummary.query.filter_by(user_id=user.id).order_by(SavedSummary.id.desc()).all()
    return render_template("history.html", conversations=convos, summaries=sums,
                           limit=FREE_MAX_CONVERSATIONS)


@app.route("/conversation/<int:cid>/share", methods=["POST"])
@login_required
def share_conversation(cid):
    convo = db.session.get(Conversation, cid)
    if not convo or convo.user_id != current_user().id:
        abort(404)
    convo.is_public = not convo.is_public
    if convo.is_public and not convo.share_token:
        convo.share_token = secrets.token_urlsafe(9)
    db.session.commit()
    flash("Public share link created." if convo.is_public else "Sharing turned off.",
          "success")
    return redirect(url_for("history"))


@app.route("/s/<token>")
def shared(token):
    convo = Conversation.query.filter_by(share_token=token, is_public=True).first()
    if not convo:
        abort(404)
    messages = [{"role": m.role, "text": m.content} for m in convo.messages]
    comments = Comment.query.filter_by(conversation_id=convo.id).order_by(Comment.id.desc()).all()
    return render_template("shared.html", convo=convo, messages=messages, comments=comments)


@app.route("/conversation/<int:cid>/delete", methods=["POST"])
@login_required
def delete_conversation(cid):
    convo = db.session.get(Conversation, cid)
    if convo and convo.user_id == current_user().id:
        db.session.delete(convo)
        db.session.commit()
        flash("Chat deleted.", "success")
    return redirect(url_for("history"))


# --------------------------------------------------------------------------- #
# Summarize / jobs / scholarships
# --------------------------------------------------------------------------- #
@app.route("/summarize", methods=["GET", "POST"])
@login_required
@limiter.limit("40 per hour")
def summarize():
    if request.method == "GET":
        return render_template("summarize.html")
    text = request.form.get("text", "")
    source = "pasted text"
    upload = request.files.get("document")
    if upload and upload.filename:
        extracted, err = extract_file_text(upload)
        if err:
            flash(err, "error")
            return render_template("summarize.html", text=text), 400
        text, source = extracted, upload.filename
    if not text.strip():
        flash("Paste some text or upload a document.", "error")
        return render_template("summarize.html"), 400
    summary, stats = summarize_text(text, max_sentences=5)
    db.session.add(SavedSummary(user_id=current_user().id, source=source, content=summary))
    db.session.commit()
    return render_template("summarize.html", summary=summary, stats=stats, source=source, saved=True)


@app.route("/jobs")
@login_required
@limiter.limit("60 per hour")
def jobs():
    q = request.args.get("q", "").strip()
    category = request.args.get("category", "").strip()
    company = request.args.get("company", "").strip()
    location = request.args.get("location", "").strip()
    results, err = (None, None)
    if q or category or company or location:
        results, err = fetch_jobs(q, category, company, location)
    return render_template("jobs.html", jobs=results, error=err, q=q, category=category,
                           company=company, location=location, categories=JOB_CATEGORIES)


@app.route("/scholarships")
@login_required
def scholarships():
    return render_template("scholarships.html", scholarships=SCHOLARSHIPS)


# --------------------------------------------------------------------------- #
# Data analysis (pandas + matplotlib)
# --------------------------------------------------------------------------- #
def _read_dataframe(fs):
    name = (fs.filename or "").lower()
    if name.endswith(".csv"):
        return pd.read_csv(fs), None
    if name.endswith((".xlsx", ".xls")):
        return pd.read_excel(fs), None
    return None, "Please upload a .csv or Excel (.xlsx) file."


def _fig_uri(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=92, bbox_inches="tight")
    _plt().close(fig)
    return "data:image/png;base64," + base64.b64encode(buf.getvalue()).decode()


def _make_charts(df):
    plt = _plt()
    charts = []
    num = df.select_dtypes("number")
    for col in list(num.columns)[:3]:
        s = df[col].dropna()
        if s.empty:
            continue
        fig, ax = plt.subplots(figsize=(4.2, 3))
        ax.hist(s, bins=20, color="#6366f1")
        ax.set_title("Distribution — " + str(col), fontsize=10)
        ax.tick_params(labelsize=7)
        charts.append({"title": "Distribution of %s" % col, "img": _fig_uri(fig)})
    for col in list(df.select_dtypes(exclude="number").columns)[:1]:
        vc = df[col].astype(str).value_counts().head(10)
        if 1 < len(vc) <= 30:
            fig, ax = plt.subplots(figsize=(4.2, 3))
            vc.plot.bar(ax=ax, color="#22d3ee")
            ax.set_title("Top values — " + str(col), fontsize=10)
            ax.tick_params(labelsize=7)
            charts.append({"title": "Value counts of %s" % col, "img": _fig_uri(fig)})
    if num.shape[1] >= 2:
        corr = num.corr(numeric_only=True)
        fig, ax = plt.subplots(figsize=(4.6, 3.6))
        im = ax.imshow(corr, cmap="viridis", vmin=-1, vmax=1)
        ax.set_xticks(range(len(corr.columns)))
        ax.set_xticklabels(corr.columns, rotation=90, fontsize=6)
        ax.set_yticks(range(len(corr.columns)))
        ax.set_yticklabels(corr.columns, fontsize=6)
        fig.colorbar(im, fraction=0.046)
        ax.set_title("Correlation", fontsize=10)
        charts.append({"title": "Correlation heatmap", "img": _fig_uri(fig)})
    return charts


@app.route("/analyze", methods=["GET", "POST"])
@pro_required
@limiter.limit("30 per hour")
def analyze():
    if request.method == "GET" or not ANALYSIS_OK:
        return render_template("analyze.html", available=ANALYSIS_OK)

    upload = request.files.get("dataset")
    if not upload or not upload.filename:
        flash("Upload a CSV or Excel file.", "error")
        return render_template("analyze.html", available=True), 400
    try:
        df, err = _read_dataframe(upload)
    except Exception as exc:  # noqa: BLE001
        flash("Couldn't read that file: %s" % exc, "error")
        return render_template("analyze.html", available=True), 400
    if err:
        flash(err, "error")
        return render_template("analyze.html", available=True), 400

    if df.shape[1] > 60:
        df = df.iloc[:, :60]
    raw_rows = len(df)

    # ---- cleaning ----
    ops = request.form.getlist("clean")
    log = []
    if "dedupe" in ops:
        before = len(df); df = df.drop_duplicates()
        log.append("Removed %d duplicate rows." % (before - len(df)))
    if "dropna" in ops:
        before = len(df); df = df.dropna()
        log.append("Dropped %d rows with missing values." % (before - len(df)))
    if "fillnum" in ops:
        num_cols = df.select_dtypes("number").columns
        df[num_cols] = df[num_cols].fillna(df[num_cols].mean(numeric_only=True))
        log.append("Filled missing numeric values with the column mean.")
    if "fillcat" in ops:
        for c in df.select_dtypes(exclude="number").columns:
            if df[c].isna().any() and not df[c].mode(dropna=True).empty:
                df[c] = df[c].fillna(df[c].mode(dropna=True)[0])
        log.append("Filled missing text values with the most frequent value.")

    # work on a sample for stats/charts if very large
    work = df.sample(20000, random_state=1) if len(df) > 20000 else df

    columns = []
    for c in df.columns:
        miss = int(df[c].isna().sum())
        columns.append({"name": str(c), "dtype": str(df[c].dtype), "missing": miss,
                        "missing_pct": round(miss * 100 / max(len(df), 1), 1),
                        "unique": int(df[c].nunique(dropna=True))})

    stats = {
        "rows": len(df), "raw_rows": raw_rows, "cols": df.shape[1],
        "missing": int(df.isna().sum().sum()),
        "duplicates": int(df.duplicated().sum()),
        "sampled": len(df) > 20000,
    }
    describe_html = ""
    num = work.select_dtypes("number")
    if not num.empty:
        describe_html = num.describe().round(3).to_html(classes="data-table", border=0)
    preview_html = df.head(8).to_html(classes="data-table", border=0, index=False)

    try:
        charts = _make_charts(work)
    except Exception as exc:  # noqa: BLE001
        app.logger.warning("chart error: %s", exc)
        charts = []

    interpretation = _interpret_dataset(work, stats, columns)

    return render_template("analyze.html", available=True, done=True,
                           filename=upload.filename, stats=stats, columns=columns,
                           describe_html=describe_html, preview_html=preview_html,
                           charts=charts, clean_log=log, interpretation=interpretation)


def _interpret_dataset(df, stats, columns):
    """Ask Gemini to interpret the profile + correlations (Markdown). None if unavailable."""
    lines = [f"Dataset: {stats['rows']} rows x {stats['cols']} columns.",
             f"Total missing cells: {stats['missing']}. Duplicate rows: {stats['duplicates']}.",
             "Columns (name: dtype, missing%, unique values):"]
    for c in columns[:40]:
        lines.append(f"- {c['name']}: {c['dtype']}, {c['missing_pct']}% missing, {c['unique']} unique")
    num = df.select_dtypes("number")
    if num.shape[1] >= 2:
        corr = num.corr(numeric_only=True)
        cols = list(corr.columns)
        pairs = []
        for i in range(len(cols)):
            for j in range(i + 1, len(cols)):
                v = corr.iloc[i, j]
                if pd.notna(v) and abs(v) >= 0.5:
                    pairs.append((abs(v), cols[i], cols[j], round(float(v), 2)))
        pairs.sort(reverse=True)
        if pairs:
            lines.append("Strong correlations:")
            for _, a1, b1, v in pairs[:6]:
                lines.append(f"- {a1} vs {b1}: {v}")
    prompt = (
        "You are a senior data analyst. Based on this dataset profile, write a "
        "concise interpretation in Markdown with three short sections using "
        "'###' headings: Data quality, Key patterns, and Suggested next steps. "
        "Use bullet points; be specific and practical.\n\n" + "\n".join(lines)
    )
    reply, err = gemini_reply([{"role": "user", "text": prompt}])
    return None if err else reply


# --------------------------------------------------------------------------- #
# Study Mode (Pro) — Gemini-generated learning material
# --------------------------------------------------------------------------- #
STUDY_MODES = {
    "flashcards": "Create 8–10 study flashcards from the material. Format each as "
                  "'**Q:** …' on one line then '**A:** …' on the next.",
    "quiz": "Create a 6-question multiple-choice quiz (options A–D). Put an "
            "**Answer key** at the very end.",
    "keypoints": "Summarise the material into concise key points as a bulleted "
                 "list, bolding the important terms.",
    "explain": "Explain the material simply, as if teaching a beginner — include "
               "a short analogy and an example.",
}


@app.route("/study", methods=["GET", "POST"])
@pro_required
@limiter.limit("30 per hour")
def study():
    if request.method == "GET":
        return render_template("study.html")
    topic = request.form.get("topic", "").strip()
    mode = request.form.get("mode", "keypoints")
    if not topic:
        flash("Enter a topic or paste your notes.", "error")
        return render_template("study.html"), 400
    instr = STUDY_MODES.get(mode, STUDY_MODES["keypoints"])
    reply, err = gemini_reply([{"role": "user", "text": f"{instr}\n\nMaterial / topic:\n{topic[:6000]}"}])
    if err:
        flash(err, "error")
        return render_template("study.html", topic=topic, mode=mode), 502
    return render_template("study.html", topic=topic, mode=mode, result=reply)


# --------------------------------------------------------------------------- #
# Productivity (Pro) — tasks, notes, transcript action-items
# --------------------------------------------------------------------------- #
@app.route("/productivity")
@pro_required
def productivity():
    uid = current_user().id
    tasks = Task.query.filter_by(user_id=uid).order_by(Task.done, Task.id.desc()).all()
    notes = Note.query.filter_by(user_id=uid).order_by(Note.id.desc()).all()
    return render_template("productivity.html", tasks=tasks, notes=notes)


@app.route("/task/add", methods=["POST"])
@pro_required
def task_add():
    text = request.form.get("text", "").strip()
    due = request.form.get("due", "").strip() or None
    if text:
        db.session.add(Task(user_id=current_user().id, text=text[:300], due=due))
        db.session.commit()
    return redirect(url_for("productivity"))


@app.route("/task/<int:tid>/toggle", methods=["POST"])
@pro_required
def task_toggle(tid):
    t = db.session.get(Task, tid)
    if t and t.user_id == current_user().id:
        t.done = not t.done
        db.session.commit()
    return redirect(url_for("productivity"))


@app.route("/task/<int:tid>/delete", methods=["POST"])
@pro_required
def task_delete(tid):
    t = db.session.get(Task, tid)
    if t and t.user_id == current_user().id:
        db.session.delete(t)
        db.session.commit()
    return redirect(url_for("productivity"))


@app.route("/note/add", methods=["POST"])
@pro_required
def note_add():
    title = request.form.get("title", "").strip()
    body = request.form.get("body", "").strip()
    if title:
        db.session.add(Note(user_id=current_user().id, title=title[:160], body=body))
        db.session.commit()
    return redirect(url_for("productivity"))


@app.route("/note/<int:nid>/delete", methods=["POST"])
@pro_required
def note_delete(nid):
    n = db.session.get(Note, nid)
    if n and n.user_id == current_user().id:
        db.session.delete(n)
        db.session.commit()
    return redirect(url_for("productivity"))


@app.route("/transcript", methods=["POST"])
@pro_required
@limiter.limit("20 per hour")
def transcript():
    text = request.form.get("transcript", "").strip()
    uid = current_user().id
    tasks = Task.query.filter_by(user_id=uid).order_by(Task.done, Task.id.desc()).all()
    notes = Note.query.filter_by(user_id=uid).order_by(Note.id.desc()).all()
    if not text:
        flash("Paste a meeting transcript first.", "error")
        return render_template("productivity.html", tasks=tasks, notes=notes), 400
    prompt = ("Summarise this meeting transcript in Markdown with three '###' "
              "sections: Summary (2–3 lines), Decisions, and Action items "
              "(each as '- [ ] owner — task'):\n\n" + text[:8000])
    reply, err = gemini_reply([{"role": "user", "text": prompt}])
    return render_template("productivity.html", tasks=tasks, notes=notes,
                           transcript_result=(None if err else reply),
                           transcript_error=err)


# --------------------------------------------------------------------------- #
# Custom workflows (Pro) — chain AI steps over text or a document
# --------------------------------------------------------------------------- #
WORKFLOW_STEPS = [
    ("summarize", "Summarize", "Summarize the following text concisely."),
    ("keypoints", "Key points", "Extract the key points as a bulleted list."),
    ("actions", "Action items", "Extract clear action items as a Markdown checklist."),
    ("simplify", "Simplify", "Rewrite the following in simple, plain language."),
    ("improve", "Improve writing", "Improve the clarity, grammar and tone; return the improved version."),
    ("translate", "Translate", "Translate the following into {lang}, preserving meaning."),
]


@app.route("/workflow", methods=["GET", "POST"])
@pro_required
@limiter.limit("15 per hour")
def workflow():
    if request.method == "GET":
        return render_template("workflow.html", steps=WORKFLOW_STEPS)

    text = request.form.get("text", "")
    source = "pasted text"
    upload = request.files.get("document")
    if upload and upload.filename:
        extracted, err = extract_file_text(upload)
        if err:
            flash(err, "error")
            return render_template("workflow.html", steps=WORKFLOW_STEPS), 400
        text, source = extracted, upload.filename
    if not text.strip():
        flash("Paste some text or upload a document to run a workflow.", "error")
        return render_template("workflow.html", steps=WORKFLOW_STEPS), 400

    selected = request.form.getlist("steps")
    if not selected:
        flash("Pick at least one step.", "error")
        return render_template("workflow.html", steps=WORKFLOW_STEPS), 400
    lang = request.form.get("lang", "French").strip() or "French"

    current = text[:8000]
    results = []
    for key, label, instr in WORKFLOW_STEPS:          # fixed, sensible order
        if key not in selected:
            continue
        prompt = (instr.format(lang=lang) if key == "translate" else instr) + "\n\n" + current
        reply, err = gemini_reply([{"role": "user", "text": prompt}])
        if err:
            flash("Workflow stopped at '%s': %s" % (label, err), "error")
            break
        current = reply
        results.append({"label": label, "output": reply})

    saved = False
    if results and request.form.get("save_note"):
        title = (request.form.get("note_title", "").strip() or ("Workflow — " + source))[:160]
        db.session.add(Note(user_id=current_user().id, title=title, body=current[:8000]))
        db.session.commit()
        saved = True

    return render_template("workflow.html", steps=WORKFLOW_STEPS, results=results,
                           source=source, saved=saved, selected=selected, lang=lang)


# --------------------------------------------------------------------------- #
# Comments on shared chats
# --------------------------------------------------------------------------- #
@app.route("/s/<token>/comment", methods=["POST"])
@limiter.limit("10 per hour")
def add_comment(token):
    convo = Conversation.query.filter_by(share_token=token, is_public=True).first()
    if not convo:
        abort(404)
    author = (request.form.get("author", "").strip() or "Guest")[:80]
    body = request.form.get("body", "").strip()[:1000]
    if body:
        db.session.add(Comment(conversation_id=convo.id, author=author, body=body))
        db.session.commit()
        flash("Comment posted.", "success")
    return redirect(url_for("shared", token=token) + "#comments")


# --------------------------------------------------------------------------- #
# Chat with your document / dataset (RAG-lite)
# --------------------------------------------------------------------------- #
DOC_SYSTEM = ("You answer questions about the user's uploaded document using ONLY "
              "the provided context. If the answer isn't in the context, say you "
              "don't know based on the document. Be concise; use Markdown.")


def _chunk_text(text, size=1200, overlap=150):
    text = re.sub(r"\s+", " ", text or "").strip()
    chunks, i = [], 0
    while i < len(text):
        chunks.append(text[i:i + size])
        i += size - overlap
    return chunks or [""]


def _retrieve(chunks, question, k=6):
    qwords = {w for w in re.findall(r"[a-z0-9']+", question.lower()) if w not in STOPWORDS}
    scored = []
    for ch in chunks:
        cw = re.findall(r"[a-z0-9']+", ch.lower())
        scored.append((sum(1 for w in cw if w in qwords), ch))
    scored.sort(key=lambda x: -x[0])
    top = [c for s, c in scored[:k] if s > 0]
    return top or chunks[:k]


@app.route("/docchat", methods=["GET", "POST"])
@login_required
@limiter.limit("30 per hour")
def docchat():
    user = current_user()
    if request.method == "POST":
        upload = request.files.get("document")
        if not upload or not upload.filename:
            flash("Choose a document or dataset to upload.", "error")
            return redirect(url_for("docchat"))
        name = upload.filename
        if name.lower().endswith((".csv", ".xlsx", ".xls")) and ANALYSIS_OK:
            try:
                df, derr = _read_dataframe(upload)
                if derr:
                    flash(derr, "error"); return redirect(url_for("docchat"))
                text = ("Dataset %s with %d rows and %d columns.\nColumns: %s\n\nPreview:\n%s"
                        % (name, df.shape[0], df.shape[1], ", ".join(map(str, df.columns)),
                           df.head(30).to_csv(index=False)))
            except Exception as exc:  # noqa: BLE001
                flash("Couldn't read that file: %s" % exc, "error")
                return redirect(url_for("docchat"))
        else:
            text, err = extract_file_text(upload)
            if err:
                flash(err, "error"); return redirect(url_for("docchat"))
        if not (text or "").strip():
            flash("Couldn't extract any text from that file.", "error")
            return redirect(url_for("docchat"))
        doc = Document(user_id=user.id, name=name[:200], text=text[:200000])
        db.session.add(doc)
        db.session.commit()
        return redirect(url_for("docchat", d=doc.id))

    docs = Document.query.filter_by(user_id=user.id).order_by(Document.id.desc()).all()
    active = None
    did = request.args.get("d", type=int)
    if did:
        active = db.session.get(Document, did)
        if not active or active.user_id != user.id:
            abort(404)
    return render_template("docchat.html", docs=docs, active=active)


@app.route("/api/docchat", methods=["POST"])
@csrf.exempt
@limiter.limit("20 per minute")
def api_docchat():
    user = current_user()
    if user is None:
        return jsonify(reply="Please log in.", error=True), 401
    payload = request.get_json(silent=True) or {}
    doc = db.session.get(Document, payload.get("doc_id") or 0)
    if not doc or doc.user_id != user.id:
        return jsonify(reply="Document not found.", error=True), 404
    question = str(payload.get("question", "")).strip()
    if not question:
        return jsonify(reply="Ask a question about the document.", error=True), 400
    context = "\n---\n".join(_retrieve(_chunk_text(doc.text), question))
    prompt = "Context from \"%s\":\n%s\n\nQuestion: %s" % (doc.name, context[:16000], question)
    reply, err = gemini_reply([{"role": "user", "text": prompt}], system=DOC_SYSTEM)
    return jsonify(reply=err or reply, error=bool(err))


# --------------------------------------------------------------------------- #
# Export (Markdown / PDF)
# --------------------------------------------------------------------------- #
def _conversation_markdown(convo):
    lines = ["# %s\n" % convo.title]
    for m in convo.messages:
        who = "**You:**" if m.role == "user" else "**Nova:**"
        lines.append("%s\n\n%s\n" % (who, m.content))
    return "\n".join(lines)


@app.route("/conversation/<int:cid>/export")
@login_required
def export_conversation(cid):
    convo = db.session.get(Conversation, cid)
    if not convo or convo.user_id != current_user().id:
        abort(404)
    fmt = request.args.get("fmt", "md")
    md_text = _conversation_markdown(convo)
    safe = re.sub(r"[^a-zA-Z0-9]+", "-", convo.title)[:40].strip("-") or "chat"

    if fmt == "pdf":
        try:                                     # lazy-load the heavy PDF stack
            import markdown as md_lib
            from xhtml2pdf import pisa
        except ImportError:
            flash("PDF export isn't available on the server.", "error")
            return redirect(url_for("history"))
        html = ("<html><head><meta charset='utf-8'><style>"
                "body{font-family:Helvetica,Arial,sans-serif;font-size:11pt;color:#222}"
                "h1{color:#4f46e5} pre{background:#f3f4fb;padding:8px;border-radius:6px}"
                "code{background:#eef0fe;padding:1px 4px}</style></head><body>"
                + md_lib.markdown(md_text, extensions=["fenced_code", "tables"])
                + "</body></html>")
        out = io.BytesIO()
        pisa.CreatePDF(html, dest=out)
        out.seek(0)
        from flask import send_file
        return send_file(out, mimetype="application/pdf", as_attachment=True,
                         download_name="%s.pdf" % safe)

    from flask import Response
    return Response(md_text, mimetype="text/markdown",
                    headers={"Content-Disposition": "attachment; filename=%s.md" % safe})


@app.route("/health")
def health():
    return jsonify(gemini_configured=bool(os.environ.get("GEMINI_API_KEY")),
                   model=GEMINI_MODEL, pdf=PdfReader is not None, docx=docx is not None,
                   analysis=ANALYSIS_OK,
                   db=app.config["SQLALCHEMY_DATABASE_URI"].split(":")[0])


if __name__ == "__main__":
    app.run(debug=os.environ.get("FLASK_DEBUG", "1") == "1", port=5004)
